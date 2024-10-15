# pip install python-docx
from docx import Document


class Nomina:
    def __init__(self):
        self.path = r"C:\Users\rodrigo.cambray\Desktop"

    def crear_documento(self, empleado):
        # Creación de documento
        recibo_nomina = Document()
        # Agregar datos de empleada o empleado
        recibo_nomina.add_heading("Recibo de nomina del empleado" + empleado[0])
        recibo_nomina.add_paragraph("Puesto: " + empleado[1])

        # Agregar datos de los pagos de empleada o empleado
        sueldo = empleado[2] * empleado[3]
        recibo_nomina.add_paragraph("Sueldo de esta semana: " + str(sueldo))

        print("Guardar documento")
        recibo_nomina.save(self.path + "/nomina.docx")
